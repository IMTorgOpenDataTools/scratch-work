#!/usr/bin/env python3
"""
Extraction functions for each file type to be used with Document class

Notes on getting text:
* assume docs are max 100 pages - only get N bytes of string to conserve memory
* `text_lst` itesms should be in batches, such as doc pages or paragraphs (don't mess up sentenceizer)
* `docs = spacy.pipe(text_lst, n_process=-1)`   #as many processes as CPUs
* resulting docs is a generator, but return as list `list(docs)`, currently
* then Document performs text processing

TODO:
* convert all file types to pdf, then extract using pdf.miner
* limit retrieval to N excerpts
* use optimal storage: `from collections import namedtuple`   #ref: https://stackoverflow.com/questions/1336791/dictionary-vs-object-which-is-more-efficient-and-why
"""

__author__ = "Jason Beach"
__version__ = "0.1.0"
__license__ = "MIT"

from backend.structures.url import UniformResourceLocator
from backend.structures.document_record import DocumentRecord
from backend.structures.document_utils import (
    MAX_PAGE_EXTRACT,
    timeout
)


import datetime
import pandas as pd

#docx
import docx

#html
import bs4
from bs4.element import Comment
from xhtml2pdf import pisa 

#pdf
import pypdf
import pdftitle                                                                 #uses pdfminer
from pdfminer.high_level import extract_text as pdf_extract_text                #uses pdfminer.six, problem TODO???
from pdfminer.high_level import extract_pages as pdf_extract_pages
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import resolve1
import fitz    
#PyMuPDF is faster than pdfminer, but not as accurate, [ref](https://medium.com/social-impact-analytics/comparing-4-methods-for-pdf-text-extraction-in-python-fd34531034f)
#TODO:PyMuPDF is similar to pikepdf but test for performance


from pathlib import Path, PosixPath




docrec = DocumentRecord()
record = docrec._asdict()


def get_title(txt):
    pass

def clean_text(txt):
    if type(txt) is list:
        combined_txt = ('.').join(txt)
        return combined_txt.replace('.','.  ').replace('\n',' ')
    elif type(txt) is str:
        txts = txt.split('.\n')
        #TODO:if the len(item)<50, then append to the earlier item
        txts = [txt.replace('-\n','').replace('\n',' ') for txt in txts]
        return txts
    else:
        return txt

def html_string_to_pdf(content, output):
    """
    Generate a pdf using a string content

    content : str - content to write in the pdf file
    output  : str - name of the file to create
    """
    # Open file to write
    result_file = open(output, "w+b") # w+b to write in binary mode.
    # convert HTML to PDF
    pisa_status = pisa.CreatePDF(
            content,                   # the HTML to convert
            dest=result_file           # file handle to recieve result
    )           
    # close output file
    result_file.close()
    result = pisa_status.err
    if not result:
        print("Successfully created PDF")
    else:
        print("Error: unable to create the PDF")    
    # return False on success and True on errors
    return result







def extract_docx_original(self, logger):
    """Extract from docx filetype.
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt = [par.text for par in excerpts.paragraphs]

    record['title'] = excerpts.paragraphs[0].text     #<<< get_title(txt)
    record['length_lines'] = len( clean_text(txt).split('.') )
    record['text'] = txt
    return record


def extract_docx(self, logger):
    """Extract from docx using `extract_pdf`
    
    'text': excerpts.paragraphs
    """
    try:
        excerpts = docx.Document(self.filepath)
    except:
        logger.info("TypeError: document not of type `.docx`")
        return None
    txt_lst = [par.text for par in excerpts.paragraphs]
    txt = ''.join(txt_lst)

    tmp_file = self.filepath.parent / (self.filepath.stem + '.pdf')
    html_string_to_pdf(content = txt, 
                       output = tmp_file
                      )
    self.filepath = tmp_file
    record = extract_pdf(self, logger)
    return record


def extract_html_original(self, logger):
    """Extract for html filetype.

    Use the optimal module, first, this is typically 
    PyMuPDF(fitz) for speed of extraction.  Then 
    try pdfminer.six, which is more accurate,  if any 
    errors occur.
    
    'text': visible text
    """
    #ref: https://stackoverflow.com/questions/1936466/how-to-scrape-only-visible-webpage-text-with-beautifulsoup
    def tag_visible(element):
        if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']:
            return False
        if isinstance(element, Comment):
            return False
        return True

    def text_from_html(soup):
        texts = soup.findAll(text=True)
        visible_texts = filter(tag_visible, texts)  
        return u" ".join(t.strip() for t in visible_texts)

    with open(self.filepath.__str__(), 'r') as f:
        html = f.readlines()
    clean_html = ('').join(html).replace('\n','')
    try:
        excerpts = bs4.BeautifulSoup(clean_html, 'html.parser')
    except:
        logger.info("TypeError: document not of type `.html`")
        return None
    txt = text_from_html(excerpts)
    record['title'] = excerpts.find('title').text
    record['text'] = txt
    return record
    

def extract_html(self, logger):
    """Extract from html using `extract_pdf`
    TODO: fix so that .pdf is removed when finished
    """
    if type(self.filepath) == UniformResourceLocator and self.filetype=='.html':
        html = self.file_str
        tmp_file = Path() / 'tmp' / (self.filename_original.split('.')[0] + '.pdf')
    elif type(self.filepath) == UniformResourceLocator and self.filetype=='.pdf':
        html = self.file_str
        tmp_file = self.filepath
    elif type(self.filepath) == PosixPath:
        with open(self.filepath.__str__(), 'r') as f:
            html = f.read()
        tmp_file = self.filepath.parent / (self.filepath.stem + '.pdf')
    else:
        raise TypeError
    html_string_to_pdf(content = html, 
                       output = tmp_file
                      )
    self.filepath = tmp_file
    record = extract_pdf(self, logger)
    return record


def extract_pdf(self, logger):
    """Extract from pdf filetype.
    
    'text': pages
    """
    def get_title(self):
        title = None
        try:
            with timeout(seconds=5):
                title = pdftitle.get_title_from_file(self.filepath.__str__())
        except Exception:
            logger.info("`pdftitle` module threw error")
            pass

        if not title:
            with timeout(seconds=5):
                with open(self.filepath.__str__(), 'rb') as f:
                    pdfReader = pypdf.PdfReader(f)
                    first_str = len(excerpts[0])
                    if pdfReader.metadata['/Title']:
                        title = pdfReader.metadata['/Title']
                    elif first_str>0 and first_str<100:
                        title = first_str
                    else:
                        pass
        return title

    def get_metadata(self):
        author = None
        if type(self.filepath) == UniformResourceLocator:
            with fitz.open(stream=self.file_str) as doc:
                tmp = doc.metadata
                tmp['page_count'] = len(doc)
        elif type(self.filepath) == PosixPath:
            with fitz.open(filename=self.filepath.__str__()) as doc:
                tmp = doc.metadata
                tmp['page_count'] = len(doc)
        else:
            raise TypeError
        result = {'author': tmp['author'], 
                  'subject': tmp['subject'], 
                  'keywords': tmp['keywords'],
                  'page_count': tmp['page_count'], 
                  'date': datetime.datetime.strptime(tmp['creationDate'].split('D:')[1][:8], "%Y%m%d").date()
                  }
        return result

    def get_toc(self):
        outlines = None
        with fitz.open(self.filepath.__str__()) as doc:
            tmp = doc.get_toc()
        outlines = tmp if tmp != [] else None
        if not outlines:
            with open(self.filepath.__str__(), 'rb') as fp:
                parser = PDFParser(fp)
                document = PDFDocument(parser)
                try:
                    outlines = list(document.get_outlines())
                except:
                    pass
        return outlines

    def get_raw_text(self, number_of_pages_to_extract_text):
        """Get raw text from pdf.

        Ensure only a limited number of pages are extracted.
        """
        MAX_TIME_SEC = 10
        raw_text = ''

        try:
            with fitz.open(self.filepath.__str__() ) as doc:
                pg_idx = 0
                for page in doc:
                    if pg_idx <= number_of_pages_to_extract_text:
                        raw_text += page.get_text()
                        pg_idx += 1
        except Exception:
            pass    

        if raw_text == '':
            try:
                with timeout(seconds=MAX_TIME_SEC):
                    raw_text = pdf_extract_text(pdf_file = self.filepath.__str__(),
                                                maxpages = number_of_pages_to_extract_text
                                                )
            except Exception:
                logger.info(f'{self.filepath.__str__()} took more than {MAX_TIME_SEC}sec to extract text')

            if not raw_text=='':
                raw_text = pdf_extract_text(pdf_file = self.filepath.__str__(),
                                            maxpages = 1
                                            )
        return raw_text        


    #process raw data
    metadata_results = get_metadata(self)
    #with fitz.open(self.filepath.__str__() ) as doc:
    #    page_count = len(doc)
    page_list_length = metadata_results['page_count'] - 1
    if MAX_PAGE_EXTRACT: 
        number_of_pages_to_extract_text = page_list_length if page_list_length <= MAX_PAGE_EXTRACT else MAX_PAGE_EXTRACT
    else:
        number_of_pages_to_extract_text = page_list_length
    raw_text = get_raw_text(self, number_of_pages_to_extract_text)
    excerpts = clean_text(raw_text)

    #create record
    #record['title'] = get_title(self)
    record['title'] = ''
    record['author'] = metadata_results['author']
    record['subject'] = metadata_results['subject']
    record['keywords'] = metadata_results['keywords']
    record['date'] = metadata_results['date']
    record['page_nos'] = metadata_results['page_count']
    #record['toc'] = get_toc(self)
    record['toc'] = ''
    record['body'] = excerpts

    return record


def extract_csv(self, logger):
    df = pd.read_csv(self.filepath.__str__(), 
                    nrows=1
                    )
    record['title'] = df.columns[0].replace('\n','')
    return record


def extract_xlsx(self, logger):
    sheets = pd.ExcelFile(self.filepath.__str__()).sheet_names
    df = pd.read_excel(self.filepath.__str__(), 
                        sheet_name=sheets[0], 
                        nrows=1
                        )

    record['title'] = df.columns[0].replace('\n','')
    return record